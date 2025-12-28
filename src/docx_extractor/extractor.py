import re
import pandas as pd
from pathlib import Path
from docx import Document

from .utils import is_section_heading, clean_numeric_like, clean_sheet_name
from .pdf_converter import PDFConverter


# --- regex helpers ---
_DIGIT_RE = re.compile(r"\d")

# Unicode superscripts: ¹²³ plus ⁰–⁹ (U+2070–U+2079)
_UNICODE_SUP_RE = re.compile(r"[\u00B9\u00B2\u00B3\u2070-\u2079]")

# Unicode subscripts: ₀–₉ (U+2080–U+2089). (There are also letter subscripts in phonetic sets,
# but for your "numeric subscripts in numbers" this is the relevant block.)
_UNICODE_SUB_RE = re.compile(r"[\u2080-\u2089]")


class DocxTableExtractor:
    def __init__(
        self,
        clean_data: bool = True,
        auto_convert_pdf: bool = True,
        small_font_ratio: float = 0.70,  # used only for superscript heuristic
    ):
        self.clean_data = clean_data
        self.auto_convert_pdf = auto_convert_pdf
        self.small_font_ratio = small_font_ratio
        self.pdf_converter = PDFConverter() if auto_convert_pdf else None

    def _cell_text_and_flags(self, cell):
        """
        Returns (cell_text, sup_flag, sub_flag)

        Superscript flag criteria:
          A) Explicit: run.font.superscript == True AND digit present in the run
          B) Unicode superscript digits present (¹²³⁰–⁹)
          C) Heuristic: numeric run font size < ratio * median cell font size

        Subscript flag criteria:
          A) Explicit: run.font.subscript == True AND digit present in the run
          B) Unicode subscript digits present (₀–₉)

        Notes:
          - No mutation of text; flags are metadata only.
          - Subscript has no font-size heuristic by default (low signal).
        """
        text = cell.text.strip()

        # Unicode markers (cheap, high precision)
        sup_unicode = bool(_UNICODE_SUP_RE.search(text))
        sub_unicode = bool(_UNICODE_SUB_RE.search(text))

        runs = []
        sizes = []

        for para in cell.paragraphs:
            for run in para.runs:
                rtxt = (run.text or "").strip()
                if not rtxt:
                    continue
                runs.append(run)
                if run.font.size is not None:
                    sizes.append(int(run.font.size))

        # Explicit flags from run properties
        sup_explicit = False
        sub_explicit = False
        for run in runs:
            rtxt = run.text or ""
            if not _DIGIT_RE.search(rtxt):
                continue
            if run.font.superscript is True:
                sup_explicit = True
            if run.font.subscript is True:
                sub_explicit = True
            if sup_explicit and sub_explicit:
                break

        # Superscript heuristic (only if we still haven't flagged superscript explicitly/unicode)
        sup_heur = False
        if not (sup_unicode or sup_explicit) and sizes:
            sizes.sort()
            n = len(sizes)
            median = sizes[n // 2] if n % 2 else (sizes[n // 2 - 1] + sizes[n // 2]) / 2
            threshold = self.small_font_ratio * median

            for run in runs:
                rtxt = run.text or ""
                if not _DIGIT_RE.search(rtxt):
                    continue
                if run.font.size is not None and int(run.font.size) < threshold:
                    sup_heur = True
                    break

        sup_flag = bool(sup_unicode or sup_explicit or sup_heur)
        sub_flag = bool(sub_unicode or sub_explicit)

        return text, sup_flag, sub_flag

    # ---------------------------------------------------------
    # Table extraction
    # ---------------------------------------------------------
    def extract_tables_by_section(self, docx_path):
        docx_path = Path(docx_path)
        if not docx_path.exists():
            raise FileNotFoundError(f"File not found: {docx_path}")

        print(f"Extracting from: {docx_path.name}")
        doc = Document(str(docx_path))

        tables_by_section = {}
        current_section = None

        for block in doc.element.body:
            if block.tag.endswith("p"):
                text = block.xpath("string(.)").strip()
                if is_section_heading(text):
                    current_section = text
                    tables_by_section.setdefault(current_section, [])

            elif block.tag.endswith("tbl"):
                table = next(t for t in doc.tables if t._element is block)

                data = []
                sup_flags = []
                sub_flags = []

                for row in table.rows:
                    row_data = []
                    row_sup = []
                    row_sub = []
                    for cell in row.cells:
                        cell_text, sup_flag, sub_flag = self._cell_text_and_flags(cell)
                        row_data.append(cell_text)
                        row_sup.append(int(sup_flag))
                        row_sub.append(int(sub_flag))
                    data.append(row_data)
                    sup_flags.append(row_sup)
                    sub_flags.append(row_sub)

                df = pd.DataFrame(data)
                sup_df = pd.DataFrame(sup_flags)
                sub_df = pd.DataFrame(sub_flags)

                if self.clean_data:
                    df = df.applymap(clean_numeric_like)

                payload = {"data": df, "sup_flags": sup_df, "sub_flags": sub_df}

                if current_section:
                    tables_by_section[current_section].append(payload)
                else:
                    tables_by_section.setdefault("UNSECTIONED", []).append(payload)

        total = sum(len(v) for v in tables_by_section.values())
        print(f"Extracted {total} tables from {len(tables_by_section)} sections")
        return tables_by_section

    # ---------------------------------------------------------
    # Excel export
    # ---------------------------------------------------------
    def export_to_excel(self, tables_by_section, output_path):
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        with pd.ExcelWriter(output_path, engine="xlsxwriter") as writer:
            for section, tables in tables_by_section.items():
                for idx, payload in enumerate(tables):
                    base = clean_sheet_name(f"{section}_{idx}")

                    payload["data"].to_excel(writer, sheet_name=base, index=False)
                    payload["sup_flags"].to_excel(
                        writer,
                        sheet_name=clean_sheet_name(f"{base}_SUP"),
                        index=False,
                    )
                    payload["sub_flags"].to_excel(
                        writer,
                        sheet_name=clean_sheet_name(f"{base}_SUB"),
                        index=False,
                    )

        print(f"Exported to {output_path.name}")

    # ---------------------------------------------------------
    # File-level processing
    # ---------------------------------------------------------
    def process_file(self, input_path, output_path, keep_intermediate=False):
        input_path = Path(input_path)
        output_path = Path(output_path)

        if input_path.suffix.lower() == ".pdf":
            if not self.auto_convert_pdf:
                raise ValueError("PDF input but auto_convert_pdf=False")
            temp_docx = input_path.with_suffix(".docx")
            self.pdf_converter.convert(input_path, temp_docx)
            docx_path = temp_docx
            cleanup = not keep_intermediate
        else:
            docx_path = input_path
            cleanup = False

        tables = self.extract_tables_by_section(docx_path)
        self.export_to_excel(tables, output_path)

        if cleanup and docx_path.exists():
            docx_path.unlink()

    # ---------------------------------------------------------
    # Batch processing
    # ---------------------------------------------------------
    def batch_process(
        self,
        input_dir,
        output_dir,
        pattern="*.pdf",
        keep_intermediate=False,
    ):
        input_dir = Path(input_dir)
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)

        files = sorted(input_dir.glob(pattern))
        print(f"Found {len(files)} files")

        for f in files:
            print(f"Processing {f.name}...")
            try:
                out = output_dir / f"{f.stem}_tables.xlsx"
                self.process_file(f, out, keep_intermediate)
            except Exception as e:
                print(f"Error processing {f.name}: {e}")

        print("Done")
